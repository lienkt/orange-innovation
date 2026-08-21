"""Shared Word-document construction helpers for the Orange Innovation Radar doc set."""
from __future__ import annotations

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor, Emu
from PIL import Image as _PILImage

ORANGE = RGBColor(0xE8, 0x6A, 0x00)
ORANGE_DARK = RGBColor(0xA8, 0x3E, 0x00)
INK = RGBColor(0x1A, 0x1A, 0x1A)
GREY = RGBColor(0x5A, 0x5A, 0x5A)
GREY_LIGHT = RGBColor(0x8A, 0x8A, 0x8A)
BLUE = RGBColor(0x2F, 0x6F, 0xB0)
GREEN = RGBColor(0x2E, 0x7D, 0x5B)
RED = RGBColor(0xA8, 0x28, 0x20)
PURPLE = RGBColor(0x6A, 0x4C, 0x93)

SH_ORANGE = "FFF1E3"
SH_GREY = "F2F2F2"
SH_HEADER = "3C3C3C"
SH_BLUE = "E6EFF8"
SH_GREEN = "E4F1EA"
SH_RED = "FBECEA"
SH_GOLD = "FAF3DF"
SH_PURPLE = "EFE9F6"

BODY_FONT = "Calibri"
MONO_FONT = "Consolas"


def _shade(el, fill):
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:color"), "auto")
    sh.set(qn("w:fill"), fill)
    el.append(sh)


def _borders(el, sides=("top", "left", "bottom", "right"), sz=4, color="D9D9D9", val="single"):
    borders = OxmlElement("w:pBdr") if el.tag.endswith("}pPr") else OxmlElement("w:tcBorders")
    for s in sides:
        e = OxmlElement(f"w:{s}")
        e.set(qn("w:val"), val)
        e.set(qn("w:sz"), str(sz))
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), color)
        borders.append(e)
    el.append(borders)


def _field(paragraph, instr):
    r = paragraph.add_run()
    fc = OxmlElement("w:fldChar"); fc.set(qn("w:fldCharType"), "begin")
    r._r.append(fc)
    r = paragraph.add_run()
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = instr
    r._r.append(it)
    r = paragraph.add_run()
    fc = OxmlElement("w:fldChar"); fc.set(qn("w:fldCharType"), "separate")
    r._r.append(fc)
    r = paragraph.add_run("…")
    r = paragraph.add_run()
    fc = OxmlElement("w:fldChar"); fc.set(qn("w:fldCharType"), "end")
    r._r.append(fc)


class Doc:
    def __init__(self, title: str, subtitle: str):
        self.d = Document()
        self.title_text = title
        self._setup_styles()
        self._setup_page(self.d.sections[0])
        self._footer(self.d.sections[0])
        self._landscape = False

    # ---------------------------------------------------------------- setup
    def _setup_page(self, sec, landscape=False):
        if landscape:
            sec.orientation = WD_ORIENT.LANDSCAPE
            sec.page_width, sec.page_height = Cm(29.7), Cm(21.0)
            sec.top_margin, sec.bottom_margin = Cm(1.5), Cm(1.5)
            sec.left_margin, sec.right_margin = Cm(1.5), Cm(1.5)
        else:
            sec.orientation = WD_ORIENT.PORTRAIT
            sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
            sec.top_margin, sec.bottom_margin = Cm(2.2), Cm(2.0)
            sec.left_margin, sec.right_margin = Cm(2.2), Cm(2.2)

    def _setup_styles(self):
        st = self.d.styles
        n = st["Normal"]
        n.font.name = BODY_FONT
        n.font.size = Pt(10.5)
        n.font.color.rgb = INK
        n.paragraph_format.space_after = Pt(7)
        n.paragraph_format.line_spacing = 1.16
        rpr = n.element.get_or_add_rPr().get_or_add_rFonts()
        rpr.set(qn("w:eastAsia"), BODY_FONT)

        for name, size, color, before, after, bold in [
            ("Heading 1", 19, ORANGE_DARK, 22, 8, True),
            ("Heading 2", 14, INK, 16, 5, True),
            ("Heading 3", 11.5, ORANGE_DARK, 12, 4, True),
            ("Heading 4", 10.5, GREY, 10, 3, True),
        ]:
            s = st[name]
            s.font.name = BODY_FONT
            s.font.size = Pt(size)
            s.font.color.rgb = color
            s.font.bold = bold
            s.paragraph_format.space_before = Pt(before)
            s.paragraph_format.space_after = Pt(after)
            s.paragraph_format.keep_with_next = True

    def _footer(self, sec):
        p = sec.footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(self.title_text + "    ·    page ")
        r.font.size = Pt(8); r.font.color.rgb = GREY_LIGHT; r.font.name = BODY_FONT
        _field(p, " PAGE ")
        for run in p.runs:
            run.font.size = Pt(8); run.font.color.rgb = GREY_LIGHT; run.font.name = BODY_FONT

    # ---------------------------------------------------------------- blocks
    def cover(self, title, subtitle, kicker, meta_rows, statement=None):
        p = self.d.add_paragraph()
        p.paragraph_format.space_before = Pt(70)
        r = p.add_run(kicker)
        r.font.size = Pt(11); r.font.bold = True; r.font.color.rgb = ORANGE; r.font.name = BODY_FONT

        p = self.d.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(title)
        r.font.size = Pt(30); r.font.bold = True; r.font.color.rgb = INK; r.font.name = BODY_FONT

        p = self.d.add_paragraph()
        p.paragraph_format.space_after = Pt(26)
        r = p.add_run(subtitle)
        r.font.size = Pt(14); r.font.color.rgb = GREY; r.font.name = BODY_FONT

        self.rule(ORANGE)

        if statement:
            p = self.d.add_paragraph()
            p.paragraph_format.space_before = Pt(16)
            p.paragraph_format.space_after = Pt(20)
            r = p.add_run(statement)
            r.font.size = Pt(11); r.font.color.rgb = INK; r.font.italic = True; r.font.name = BODY_FONT

        t = self.d.add_table(rows=0, cols=2)
        t.alignment = WD_TABLE_ALIGNMENT.LEFT
        t.autofit = False
        for k, v in meta_rows:
            row = t.add_row()
            c0, c1 = row.cells
            c0.width = Cm(3.6); c1.width = Cm(13.0)
            r = c0.paragraphs[0].add_run(k)
            r.font.size = Pt(9); r.font.bold = True; r.font.color.rgb = GREY; r.font.name = BODY_FONT
            r = c1.paragraphs[0].add_run(v)
            r.font.size = Pt(9); r.font.color.rgb = INK; r.font.name = BODY_FONT
            for c in (c0, c1):
                c.paragraphs[0].paragraph_format.space_after = Pt(2)
                _borders(c._tc.get_or_add_tcPr(), ("top", "left", "bottom", "right"), sz=0, val="none")
        self.pagebreak()

    def rule(self, color=ORANGE, sz=12):
        p = self.d.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        pPr = p._p.get_or_add_pPr()
        bd = OxmlElement("w:pBdr")
        b = OxmlElement("w:bottom")
        b.set(qn("w:val"), "single"); b.set(qn("w:sz"), str(sz))
        b.set(qn("w:space"), "1")
        b.set(qn("w:color"), "%02X%02X%02X" % (color[0], color[1], color[2]) if isinstance(color, tuple) else str(color))
        bd.append(b); pPr.append(bd)
        return p

    def toc(self, entries):
        self._ensure_portrait()
        self.h1("Contents", numbered=False)
        for level, text in entries:
            p = self.d.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.left_indent = Cm(0.0 if level == 1 else 0.8)
            r = p.add_run(text)
            r.font.name = BODY_FONT
            r.font.size = Pt(10.5 if level == 1 else 9.5)
            r.font.bold = level == 1
            r.font.color.rgb = INK if level == 1 else GREY
        self.pagebreak()

    def h1(self, text, numbered=True):
        self._ensure_portrait()
        h = self.d.add_heading(text, level=1)
        for r in h.runs:
            r.font.name = BODY_FONT
        return h

    def h2(self, text):
        self._ensure_portrait()
        h = self.d.add_heading(text, level=2)
        for r in h.runs:
            r.font.name = BODY_FONT
        return h

    def h3(self, text):
        self._ensure_portrait()
        h = self.d.add_heading(text, level=3)
        for r in h.runs:
            r.font.name = BODY_FONT
        return h

    def p(self, text, size=10.5, color=INK, italic=False, bold=False, after=7, before=0, indent=0):
        self._ensure_portrait()
        par = self.d.add_paragraph()
        par.paragraph_format.space_after = Pt(after)
        par.paragraph_format.space_before = Pt(before)
        if indent:
            par.paragraph_format.left_indent = Cm(indent)
        self._rich(par, text, size, color, italic, bold)
        return par

    def _rich(self, par, text, size, color, italic, bold):
        """**bold** and `mono` inline markers."""
        import re
        for chunk in re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text):
            if not chunk:
                continue
            r = par.add_run()
            if chunk.startswith("**") and chunk.endswith("**"):
                r.text = chunk[2:-2]; r.font.bold = True; r.font.name = BODY_FONT
            elif chunk.startswith("`") and chunk.endswith("`"):
                r.text = chunk[1:-1]; r.font.name = MONO_FONT; r.font.size = Pt(size - 1.0)
                r.font.color.rgb = ORANGE_DARK
            else:
                r.text = chunk; r.font.name = BODY_FONT
                r.font.bold = bold
            if r.font.size is None:
                r.font.size = Pt(size)
            if r.font.color.rgb is None:
                r.font.color.rgb = color
            r.font.italic = italic

    def bullets(self, items, size=10.5, style="List Bullet", after=3):
        self._ensure_portrait()
        for it in items:
            par = self.d.add_paragraph(style=style)
            par.paragraph_format.space_after = Pt(after)
            par.paragraph_format.left_indent = Cm(0.7)
            self._rich(par, it, size, INK, False, False)

    def numbers(self, items, size=10.5):
        self.bullets(items, size=size, style="List Number")

    def callout(self, title, body, fill=SH_ORANGE, accent=ORANGE_DARK):
        self._ensure_portrait()
        t = self.d.add_table(rows=1, cols=1)
        cell = t.cell(0, 0)
        _shade(cell._tc.get_or_add_tcPr(), fill)
        _borders(cell._tc.get_or_add_tcPr(), ("top", "bottom", "right"), sz=0, val="none")
        _borders(cell._tc.get_or_add_tcPr(), ("left",), sz=18, color="%02X%02X%02X" % (accent[0], accent[1], accent[2]))
        cell.paragraphs[0]._p.getparent().remove(cell.paragraphs[0]._p)
        if title:
            p = cell.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
            r = p.add_run(title)
            r.font.size = Pt(10); r.font.bold = True; r.font.color.rgb = accent; r.font.name = BODY_FONT
        for i, line in enumerate(body if isinstance(body, list) else [body]):
            p = cell.add_paragraph()
            p.paragraph_format.space_after = Pt(2 if i < len(body) - 1 else 0)
            self._rich(p, line, 9.5, INK, False, False)
        self.d.add_paragraph().paragraph_format.space_after = Pt(2)

    def table(self, headers, rows, widths=None, size=9, header_fill=SH_HEADER, zebra=True, first_bold=False):
        self._ensure_portrait()
        t = self.d.add_table(rows=1, cols=len(headers))
        t.alignment = WD_TABLE_ALIGNMENT.LEFT
        hdr = t.rows[0]
        for i, h in enumerate(headers):
            c = hdr.cells[i]
            _shade(c._tc.get_or_add_tcPr(), header_fill)
            c.paragraphs[0].paragraph_format.space_after = Pt(1)
            c.paragraphs[0].paragraph_format.space_before = Pt(1)
            r = c.paragraphs[0].add_run(h)
            r.font.size = Pt(size); r.font.bold = True
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); r.font.name = BODY_FONT
        for j, row in enumerate(rows):
            cells = t.add_row().cells
            for i, val in enumerate(row):
                c = cells[i]
                if zebra and j % 2 == 1:
                    _shade(c._tc.get_or_add_tcPr(), SH_GREY)
                pr = c.paragraphs[0]
                pr.paragraph_format.space_after = Pt(1)
                pr.paragraph_format.space_before = Pt(1)
                self._rich(pr, str(val), size, INK, False, first_bold and i == 0)
        for row in t.rows:
            for i, c in enumerate(row.cells):
                _borders(c._tc.get_or_add_tcPr(), ("top", "left", "bottom", "right"), sz=2, color="C8C8C8")
        if widths:
            t.autofit = False
            for row in t.rows:
                for i, w in enumerate(widths):
                    row.cells[i].width = Cm(w)
        self.d.add_paragraph().paragraph_format.space_after = Pt(3)
        return t

    def code(self, text, size=8.5):
        self._ensure_portrait()
        t = self.d.add_table(rows=1, cols=1)
        cell = t.cell(0, 0)
        _shade(cell._tc.get_or_add_tcPr(), "F6F6F6")
        _borders(cell._tc.get_or_add_tcPr(), ("top", "left", "bottom", "right"), sz=2, color="DDDDDD")
        cell.paragraphs[0]._p.getparent().remove(cell.paragraphs[0]._p)
        for line in text.split("\n"):
            p = cell.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            r = p.add_run(line if line else " ")
            r.font.name = MONO_FONT; r.font.size = Pt(size); r.font.color.rgb = INK
        self.d.add_paragraph().paragraph_format.space_after = Pt(3)

    def _ensure_portrait(self):
        if getattr(self, "_landscape", False):
            s = self.d.add_section(WD_SECTION.NEW_PAGE)
            self._setup_page(s, landscape=False)
            self._landscape = False

    def _ensure_landscape(self):
        if not getattr(self, "_landscape", False):
            s = self.d.add_section(WD_SECTION.NEW_PAGE)
            self._setup_page(s, landscape=True)
            self._landscape = True
            return False
        return True   # already landscape: caller must break the page itself

    def figure(self, path, caption, note=None):
        """Figures live on landscape pages; consecutive figures share one section."""
        needs_break = self._ensure_landscape()
        p = self.d.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        if needs_break:
            p.paragraph_format.page_break_before = True
        with _PILImage.open(path) as im:
            aspect = im.width / im.height
        # Landscape A4 with 1.5 cm margins: 26.7 cm usable width, ~17.2 cm usable
        # height once the caption is allowed for. Fit to whichever bound binds first.
        width_cm = min(26.4, 17.2 * aspect)
        p.add_run().add_picture(path, width=Cm(width_cm))
        if note:
            cp = self.d.add_paragraph()
            cp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            cp.paragraph_format.space_after = Pt(0)
            r = cp.add_run(note)
            r.font.size = Pt(8.5); r.font.color.rgb = GREY; r.font.italic = True; r.font.name = BODY_FONT

    def pagebreak(self):
        self.d.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    def save(self, path):
        self._ensure_portrait()
        self.d.save(path)
        print("wrote", path)
