"""Speaker-notes document construction — optimised for reading aloud."""
from __future__ import annotations
import re
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ORANGE = RGBColor(0xE8, 0x6A, 0x00)
ORANGE_D = RGBColor(0xA8, 0x3E, 0x00)
INK = RGBColor(0x14, 0x14, 0x14)
GREY = RGBColor(0x5A, 0x5A, 0x5A)
GREY_L = RGBColor(0x8E, 0x8E, 0x8E)
BLUE = RGBColor(0x2F, 0x6F, 0xB0)
GREEN = RGBColor(0x2E, 0x7D, 0x5B)
RED = RGBColor(0xA8, 0x28, 0x20)
PURPLE = RGBColor(0x6A, 0x4C, 0x93)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

FONT = "Calibri"
MONO = "Consolas"

SECTION_COLOR = {
    "OPENING": ORANGE_D, "WHY": ORANGE_D, "CONCEPT": BLUE, "STATUS": GREEN,
    "FUNCTIONALITY": PURPLE, "ARCHITECTURE": RGBColor(0x1F, 0x7A, 0x85), "CLOSE": ORANGE_D,
}


def _shade(el, fill):
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear"); sh.set(qn("w:color"), "auto"); sh.set(qn("w:fill"), fill)
    el.append(sh)


def _nobord(cell):
    b = OxmlElement("w:tcBorders")
    for s in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{s}"); e.set(qn("w:val"), "none"); b.append(e)
    cell._tc.get_or_add_tcPr().append(b)


def _hex(c):
    return "%02X%02X%02X" % (c[0], c[1], c[2])


def _field(paragraph, instr):
    r = paragraph.add_run(); fc = OxmlElement("w:fldChar"); fc.set(qn("w:fldCharType"), "begin"); r._r.append(fc)
    r = paragraph.add_run(); it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = instr; r._r.append(it)
    r = paragraph.add_run(); fc = OxmlElement("w:fldChar"); fc.set(qn("w:fldCharType"), "separate"); r._r.append(fc)
    paragraph.add_run("1")
    r = paragraph.add_run(); fc = OxmlElement("w:fldChar"); fc.set(qn("w:fldCharType"), "end"); r._r.append(fc)


class Notes:
    def __init__(self, footer_text):
        self.d = Document()
        s = self.d.sections[0]
        s.page_width, s.page_height = Cm(21.0), Cm(29.7)
        s.top_margin, s.bottom_margin = Cm(1.6), Cm(1.4)
        s.left_margin, s.right_margin = Cm(1.9), Cm(1.9)
        n = self.d.styles["Normal"]
        n.font.name = FONT; n.font.size = Pt(12); n.font.color.rgb = INK
        n.paragraph_format.space_after = Pt(8)
        n.paragraph_format.line_spacing = 1.28
        n.element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), FONT)
        p = s.footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(footer_text + "    ·    ")
        _field(p, " PAGE ")
        for run in p.runs:
            run.font.size = Pt(8); run.font.color.rgb = GREY_L; run.font.name = FONT

    # ---------------------------------------------------------------- pieces
    def rich(self, par, text, size=12, color=INK, italic=False):
        """**bold**  ·  [stage direction]  ·  «pause»"""
        for chunk in re.split(r"(\*\*[^*]+\*\*|\[[^\]]+\])", text):
            if not chunk:
                continue
            r = par.add_run()
            if chunk.startswith("**"):
                r.text = chunk[2:-2]; r.font.bold = True
            elif chunk.startswith("["):
                r.text = chunk; r.font.italic = True; r.font.color.rgb = BLUE; r.font.size = Pt(size - 1.5)
            else:
                r.text = chunk; r.font.italic = italic
            r.font.name = FONT
            if r.font.size is None:
                r.font.size = Pt(size)
            if r.font.color.rgb is None:
                r.font.color.rgb = color
        return par

    def banner(self, n, total, section, title, clock, dur):
        t = self.d.add_table(rows=1, cols=2)
        t.autofit = False
        col = SECTION_COLOR.get(section, ORANGE_D)
        left, right = t.rows[0].cells
        left.width, right.width = Cm(11.6), Cm(5.6)
        for c in (left, right):
            _shade(c._tc.get_or_add_tcPr(), _hex(col)); _nobord(c)
            c.paragraphs[0].paragraph_format.space_after = Pt(0)
            c.paragraphs[0].paragraph_format.space_before = Pt(3)
        r = left.paragraphs[0].add_run(f"SLIDE {n}  /  {total}")
        r.font.size = Pt(15); r.font.bold = True; r.font.color.rgb = WHITE; r.font.name = FONT
        r = left.paragraphs[0].add_run(f"      {section}")
        r.font.size = Pt(9.5); r.font.bold = True; r.font.color.rgb = RGBColor(0xFF, 0xE2, 0xC8); r.font.name = FONT
        p2 = right.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p2.add_run(f"{clock}    ·    {dur}")
        r.font.size = Pt(10.5); r.font.bold = True; r.font.color.rgb = WHITE; r.font.name = FONT
        sp = self.d.add_paragraph(); sp.paragraph_format.space_after = Pt(0); sp.paragraph_format.space_before = Pt(0)
        for run in sp.runs:
            run.font.size = Pt(2)

        p = self.d.add_paragraph()
        p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(1)
        r = p.add_run(title)
        r.font.size = Pt(20); r.font.bold = True; r.font.color.rgb = INK; r.font.name = FONT

    def onscreen(self, text):
        p = self.d.add_paragraph()
        p.paragraph_format.space_after = Pt(11)
        r = p.add_run("ON SCREEN   ")
        r.font.size = Pt(8.5); r.font.bold = True; r.font.color.rgb = ORANGE; r.font.name = FONT
        r = p.add_run(text)
        r.font.size = Pt(10); r.font.color.rgb = GREY; r.font.italic = True; r.font.name = FONT

    def say_label(self):
        p = self.d.add_paragraph()
        p.paragraph_format.space_after = Pt(4); p.paragraph_format.space_before = Pt(2)
        r = p.add_run("SAY")
        r.font.size = Pt(9); r.font.bold = True; r.font.color.rgb = ORANGE_D; r.font.name = FONT
        pPr = p._p.get_or_add_pPr()
        bd = OxmlElement("w:pBdr"); b = OxmlElement("w:bottom")
        b.set(qn("w:val"), "single"); b.set(qn("w:sz"), "6"); b.set(qn("w:space"), "2"); b.set(qn("w:color"), _hex(ORANGE))
        bd.append(b); pPr.append(bd)

    def say(self, paras, size=12.5):
        self.say_label()
        for t in paras:
            p = self.d.add_paragraph()
            p.paragraph_format.space_after = Pt(9)
            p.paragraph_format.line_spacing = 1.32
            self.rich(p, t, size=size)

    def box(self, label, items, accent=BLUE, fill="EAF1F9", size=10.5):
        t = self.d.add_table(rows=1, cols=1)
        c = t.cell(0, 0)
        _shade(c._tc.get_or_add_tcPr(), fill)
        b = OxmlElement("w:tcBorders")
        for s in ("top", "bottom", "right"):
            e = OxmlElement(f"w:{s}"); e.set(qn("w:val"), "none"); b.append(e)
        e = OxmlElement("w:left"); e.set(qn("w:val"), "single"); e.set(qn("w:sz"), "18")
        e.set(qn("w:space"), "0"); e.set(qn("w:color"), _hex(accent)); b.append(e)
        c._tc.get_or_add_tcPr().append(b)
        c.paragraphs[0]._p.getparent().remove(c.paragraphs[0]._p)
        p = c.add_paragraph(); p.paragraph_format.space_after = Pt(3)
        r = p.add_run(label)
        r.font.size = Pt(8.5); r.font.bold = True; r.font.color.rgb = accent; r.font.name = FONT
        for i, it in enumerate(items):
            p = c.add_paragraph()
            p.paragraph_format.space_after = Pt(2 if i < len(items) - 1 else 0)
            p.paragraph_format.line_spacing = 1.2
            self.rich(p, it, size=size)
        sp = self.d.add_paragraph(); sp.paragraph_format.space_after = Pt(4)

    def advance(self, text, terminal=False):
        t = self.d.add_table(rows=1, cols=1)
        c = t.cell(0, 0)
        col = GREEN if not terminal else ORANGE_D
        _shade(c._tc.get_or_add_tcPr(), _hex(col)); _nobord(c)
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)
        r = p.add_run(text)
        r.font.size = Pt(13); r.font.bold = True; r.font.color.rgb = WHITE; r.font.name = FONT

    def spacer(self, pts=6):
        p = self.d.add_paragraph(); p.paragraph_format.space_after = Pt(pts)

    def pagebreak(self):
        self.d.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    def h(self, text, size=17, color=ORANGE_D, before=0, after=6):
        p = self.d.add_paragraph()
        p.paragraph_format.space_before = Pt(before); p.paragraph_format.space_after = Pt(after)
        r = p.add_run(text)
        r.font.size = Pt(size); r.font.bold = True; r.font.color.rgb = color; r.font.name = FONT

    def p(self, text, size=11, color=INK, after=7, italic=False):
        p = self.d.add_paragraph()
        p.paragraph_format.space_after = Pt(after)
        self.rich(p, text, size=size, color=color, italic=italic)

    def table(self, headers, rows, widths, size=9.5, highlight=None):
        t = self.d.add_table(rows=1, cols=len(headers)); t.autofit = False
        t.alignment = WD_TABLE_ALIGNMENT.LEFT
        for i, hh in enumerate(headers):
            c = t.rows[0].cells[i]
            _shade(c._tc.get_or_add_tcPr(), "3C3C3C")
            c.paragraphs[0].paragraph_format.space_after = Pt(1)
            c.paragraphs[0].paragraph_format.space_before = Pt(1)
            r = c.paragraphs[0].add_run(hh)
            r.font.size = Pt(size); r.font.bold = True; r.font.color.rgb = WHITE; r.font.name = FONT
        for j, row in enumerate(rows):
            cells = t.add_row().cells
            for i, v in enumerate(row):
                c = cells[i]
                if highlight and highlight(j):
                    _shade(c._tc.get_or_add_tcPr(), "FFF1E3")
                elif j % 2 == 1:
                    _shade(c._tc.get_or_add_tcPr(), "F4F4F4")
                pr = c.paragraphs[0]
                pr.paragraph_format.space_after = Pt(1); pr.paragraph_format.space_before = Pt(1)
                pr.paragraph_format.line_spacing = 1.1
                self.rich(pr, str(v), size=size)
        for row in t.rows:
            for i, c in enumerate(row.cells):
                c.width = Cm(widths[i])
                b = OxmlElement("w:tcBorders")
                for s in ("top", "left", "bottom", "right"):
                    e = OxmlElement(f"w:{s}"); e.set(qn("w:val"), "single"); e.set(qn("w:sz"), "2")
                    e.set(qn("w:space"), "0"); e.set(qn("w:color"), "C8C8C8"); b.append(e)
                c._tc.get_or_add_tcPr().append(b)
        self.spacer(4)

    def kv(self, rows, w0=5.0, w1=12.2, size=10):
        t = self.d.add_table(rows=0, cols=2); t.autofit = False
        for a, b in rows:
            c0, c1 = t.add_row().cells
            c0.width, c1.width = Cm(w0), Cm(w1)
            for c in (c0, c1):
                _nobord(c)
                c.paragraphs[0].paragraph_format.space_after = Pt(3)
            r = c0.paragraphs[0].add_run(a)
            r.font.size = Pt(size); r.font.bold = True; r.font.color.rgb = GREY; r.font.name = FONT
            r = c1.paragraphs[0].add_run(b)
            r.font.size = Pt(size); r.font.color.rgb = INK; r.font.name = FONT
        self.spacer(4)

    def save(self, path):
        self.d.save(path)
        print("wrote", path)
